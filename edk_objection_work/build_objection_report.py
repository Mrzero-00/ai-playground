from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


REFERENCE = Path("/Users/sonsang-il/Desktop/ai/edk_objection_work/EDK_renewal_completion_report_v0.1.6_draft.docx")
OUTPUT = Path("/Users/sonsang-il/Desktop/ai/EDK_리뉴얼_프로젝트_이의제기_및_검토의견서_v1.0.docx")

NAVY = "17365D"
MID_BLUE = "2F5597"
LIGHT_BLUE = "D9E2F3"
PALE_BLUE = "EAF0F8"
LIGHT_GRAY = "F2F2F2"
MID_GRAY = "A6A6A6"
DARK = "222222"
RED = "9C0006"
FONT = "Noto Sans KR"
BULLET_NUM_ID = None
NUMBER_NUM_ID = None
NUMBER_ABSTRACT_ID = None


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color="B7B7B7", size="6"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = qn(f"w:{edge}")
        node = borders.find(tag)
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), size)
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_run_font(run, size=10.5, bold=False, color=DARK, italic=False):
    run.font.name = FONT
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
        r_fonts.set(qn(f"w:{attr}"), FONT)


def get_style(doc, name):
    for style in doc.styles:
        if style.name == name:
            return style
    return None


def style_paragraph(p, before=0, after=6, line=1.35, align=None, keep=False):
    fmt = p.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line
    if align is not None:
        p.alignment = align
    if keep:
        fmt.keep_with_next = True


def add_text(doc, text, *, bold=False, color=DARK, size=10.5, before=0, after=6,
             align=None, italic=False, keep=False):
    p = doc.add_paragraph()
    style_paragraph(p, before=before, after=after, align=align, keep=keep)
    r = p.add_run(text)
    set_run_font(r, size=size, bold=bold, color=color, italic=italic)
    return p


def add_mixed(doc, runs, *, before=0, after=6, align=None, keep=False):
    p = doc.add_paragraph()
    style_paragraph(p, before=before, after=after, align=align, keep=keep)
    for item in runs:
        if isinstance(item, str):
            text, bold, color, italic = item, False, DARK, False
        else:
            text = item[0]
            bold = item[1] if len(item) > 1 else False
            color = item[2] if len(item) > 2 else DARK
            italic = item[3] if len(item) > 3 else False
        r = p.add_run(text)
        set_run_font(r, bold=bold, color=color, italic=italic)
    return p


def add_heading(doc, text, level=1):
    sizes = {1: 20, 2: 15, 3: 12}
    before = {1: 10, 2: 9, 3: 7}
    p = doc.add_paragraph()
    p.style = get_style(doc, f"Heading {level}")
    style_paragraph(p, before=before[level], after=5, line=1.1, keep=True)
    r = p.add_run(text)
    set_run_font(r, size=sizes[level], bold=True, color=NAVY)
    p.paragraph_format.keep_with_next = True
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style=get_style(doc, "List Paragraph"))
    style_paragraph(p, after=3)
    p_pr = p._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), str(level))
    num_id = OxmlElement("w:numId")
    num_id.set(qn("w:val"), str(BULLET_NUM_ID))
    num_pr.append(ilvl)
    num_pr.append(num_id)
    p_pr.append(num_pr)
    r = p.add_run(text)
    set_run_font(r)
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style=get_style(doc, "List Paragraph"))
    style_paragraph(p, after=3)
    p_pr = p._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id = OxmlElement("w:numId")
    num_id.set(qn("w:val"), str(NUMBER_NUM_ID))
    num_pr.append(ilvl)
    num_pr.append(num_id)
    p_pr.append(num_pr)
    r = p.add_run(text)
    set_run_font(r)
    return p


def add_callout(doc, label, body, fill=PALE_BLUE, accent=NAVY):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Cm(16)
    cell = table.cell(0, 0)
    cell.width = Cm(16)
    set_cell_shading(cell, fill)
    set_cell_margins(cell, top=160, start=180, bottom=150, end=180)
    set_table_borders(table, color=accent, size="8")
    p = cell.paragraphs[0]
    style_paragraph(p, after=4)
    r = p.add_run(label)
    set_run_font(r, size=11, bold=True, color=accent)
    p2 = cell.add_paragraph()
    style_paragraph(p2, after=0)
    r2 = p2.add_run(body)
    set_run_font(r2, size=10.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return table


def add_table(doc, headers, rows, widths_cm):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for idx, width in enumerate(widths_cm):
        table.columns[idx].width = Cm(width)
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for i, text in enumerate(headers):
        cell = hdr.cells[i]
        cell.width = Cm(widths_cm[i])
        set_cell_shading(cell, LIGHT_BLUE)
        set_cell_margins(cell, top=120, start=120, bottom=120, end=120)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        style_paragraph(p, after=0, line=1.15)
        r = p.add_run(text)
        set_run_font(r, size=9.5, bold=True, color=NAVY)
    for row_data in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row_data):
            cells[i].width = Cm(widths_cm[i])
            set_cell_margins(cells[i], top=105, start=120, bottom=105, end=120)
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cells[i].paragraphs[0]
            style_paragraph(p, after=0, line=1.18)
            r = p.add_run(str(value))
            set_run_font(r, size=9.2, bold=(i == 0))
    set_table_borders(table)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return table


def add_page_break(doc):
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def clear_body(doc):
    body = doc._element.body
    sect_pr = body.sectPr
    for child in list(body):
        if child is not sect_pr:
            body.remove(child)


def configure_styles(doc):
    names = {s.name for s in doc.styles}
    normal = get_style(doc, "Normal") if "Normal" in names else doc.styles.add_style("Normal", WD_STYLE_TYPE.PARAGRAPH)
    normal.font.name = FONT
    normal.font.size = Pt(10.5)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    for style_name in ("Heading 1", "Heading 2", "Heading 3", "List Paragraph"):
        style = get_style(doc, style_name)
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)


def setup_numbering(doc):
    numbering = doc.part.numbering_part.element
    abstract_ids = [int(x.get(qn("w:abstractNumId"))) for x in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(x.get(qn("w:numId"))) for x in numbering.findall(qn("w:num"))]
    next_abs = max(abstract_ids, default=0) + 1
    next_num = max(num_ids, default=0) + 1

    def make_abstract(abs_id, fmt, text, font=None):
        abstract = OxmlElement("w:abstractNum")
        abstract.set(qn("w:abstractNumId"), str(abs_id))
        multi = OxmlElement("w:multiLevelType")
        multi.set(qn("w:val"), "multilevel")
        abstract.append(multi)
        for level in range(2):
            lvl = OxmlElement("w:lvl")
            lvl.set(qn("w:ilvl"), str(level))
            start = OxmlElement("w:start")
            start.set(qn("w:val"), "1")
            num_fmt = OxmlElement("w:numFmt")
            num_fmt.set(qn("w:val"), fmt)
            lvl_text = OxmlElement("w:lvlText")
            lvl_text.set(qn("w:val"), text if level == 0 else "◦")
            suff = OxmlElement("w:suff")
            suff.set(qn("w:val"), "tab")
            p_pr = OxmlElement("w:pPr")
            tabs = OxmlElement("w:tabs")
            tab = OxmlElement("w:tab")
            tab.set(qn("w:val"), "num")
            tab.set(qn("w:pos"), str(720 + level * 360))
            tabs.append(tab)
            ind = OxmlElement("w:ind")
            ind.set(qn("w:left"), str(720 + level * 360))
            ind.set(qn("w:hanging"), "360")
            p_pr.append(tabs)
            p_pr.append(ind)
            lvl.extend([start, num_fmt, lvl_text, suff, p_pr])
            if font:
                r_pr = OxmlElement("w:rPr")
                r_fonts = OxmlElement("w:rFonts")
                r_fonts.set(qn("w:ascii"), font)
                r_fonts.set(qn("w:hAnsi"), font)
                r_pr.append(r_fonts)
                lvl.append(r_pr)
            abstract.append(lvl)
        numbering.append(abstract)
        num = OxmlElement("w:num")
        num.set(qn("w:numId"), str(next_num + (abs_id - next_abs)))
        abs_ref = OxmlElement("w:abstractNumId")
        abs_ref.set(qn("w:val"), str(abs_id))
        num.append(abs_ref)
        numbering.append(num)
        return next_num + (abs_id - next_abs)

    bullet_id = make_abstract(next_abs, "bullet", "•", "Arial")
    number_id = make_abstract(next_abs + 1, "decimal", "%1.")
    return bullet_id, number_id, next_abs + 1


def restart_numbering(doc):
    global NUMBER_NUM_ID
    numbering = doc.part.numbering_part.element
    num_ids = [int(x.get(qn("w:numId"))) for x in numbering.findall(qn("w:num"))]
    new_id = max(num_ids, default=0) + 1
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(new_id))
    abs_ref = OxmlElement("w:abstractNumId")
    abs_ref.set(qn("w:val"), str(NUMBER_ABSTRACT_ID))
    num.append(abs_ref)
    override = OxmlElement("w:lvlOverride")
    override.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:startOverride")
    start.set(qn("w:val"), "1")
    override.append(start)
    num.append(override)
    numbering.append(num)
    NUMBER_NUM_ID = new_id


def add_footer(section):
    footer = section.footer
    p = footer.paragraphs[0]
    p.clear()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    style_paragraph(p, before=3, after=0, line=1.0)
    r = p.add_run("EDK 리뉴얼 프로젝트 이의제기 및 검토의견서")
    set_run_font(r, size=8.5, color="666666")
    r2 = p.add_run(" " * 10 + "|  ")
    set_run_font(r2, size=8.5, color="999999")
    r3 = p.add_run("Page ")
    set_run_font(r3, size=8.5, color="666666")
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    p._p.append(fld)


doc = Document(REFERENCE)
clear_body(doc)
configure_styles(doc)
global_ids = setup_numbering(doc)
BULLET_NUM_ID, NUMBER_NUM_ID, NUMBER_ABSTRACT_ID = global_ids

section = doc.sections[0]
section.page_width = Cm(21.0)
section.page_height = Cm(29.7)
section.top_margin = Cm(2.5)
section.bottom_margin = Cm(2.5)
section.left_margin = Cm(2.5)
section.right_margin = Cm(2.5)
add_footer(section)

# Cover
p = doc.add_paragraph()
style_paragraph(p, before=85, after=12, align=WD_ALIGN_PARAGRAPH.CENTER)
r = p.add_run("EDK")
set_run_font(r, size=34, bold=True, color=NAVY)

p = doc.add_paragraph()
style_paragraph(p, after=8, align=WD_ALIGN_PARAGRAPH.CENTER)
r = p.add_run("리뉴얼 프로젝트 이의제기 및 검토의견서")
set_run_font(r, size=24, bold=True, color=NAVY)

p = doc.add_paragraph()
style_paragraph(p, after=36, align=WD_ALIGN_PARAGRAPH.CENTER)
r = p.add_run("UI/UX 동일성, 프론트엔드 코드 보존 및 성능개선 범위에 관한 공식 검토")
set_run_font(r, size=12, color="555555")

add_table(
    doc,
    ["항목", "내용"],
    [
        ("참조 문서", "EDK 리뉴얼 프로젝트 종료 보고서 v0.1.6 draft"),
        ("문서 성격", "종료보고서 기재 내용에 대한 이의제기 및 사실·범위 확인 요청"),
        ("작성 목적", "요구사항 해석, 범위 변경 절차, 성능 검증 범위의 명확화"),
        ("작성일", "2026-07-23"),
        ("작성", "[작성자 기입]"),
        ("수신", "[수신자 기입]"),
        ("배포 구분", "검토용"),
    ],
    [4.2, 11.8],
)

add_text(
    doc,
    "본 문서는 제공된 종료보고서 초안을 근거로 작성한 검토 의견이며, 계약서·제안요청서·회의록·변경승인서 원문이 추가되면 그 내용에 따라 보완되어야 한다.",
    size=9, color="666666", italic=True, before=18, align=WD_ALIGN_PARAGRAPH.CENTER,
)

add_page_break(doc)

# 1
add_heading(doc, "1. 요약 및 공식 이의제기", 1)
add_heading(doc, "1.1 검토 결론", 2)
add_callout(
    doc,
    "핵심 의견",
    "“UI/UX를 동일하게 유지한다”는 요구는 사용자가 보는 화면과 업무 흐름의 동등성을 뜻하며, 그 자체로 “프론트엔드 소스 코드를 일절 수정하지 않는다”는 구현 제약을 의미하지 않는다. "
    "프론트엔드와 백엔드의 V1→V2 리뉴얼 및 성능개선이 당초 범위였다면, 기존 프론트엔드를 그대로 보존하는 결정은 기술적 세부사항이 아니라 범위·산출물·검수 기준에 영향을 주는 전략 변경으로 보아야 한다.",
)
add_text(
    doc,
    "제공된 종료보고서는 한편으로 프로젝트 범위를 “백엔드 시스템의 전면 재작성과 운영 데이터 이관”으로 규정하면서, 다른 한편으로 초기에는 백엔드와 프론트엔드를 모두 신규 기술로 재작성하였고 이후 2026년 4월 25일 프론트엔드 신규 개발을 중단했다고 서술한다. "
    "따라서 쟁점은 ‘기존 프론트 코드를 사용하는 것이 기술적으로 가능한가’가 아니라, 그 전략 변경이 당초 계약 범위와 발주자의 명시적 승인에 부합하는지 여부이다."
)
add_heading(doc, "1.2 이의제기 사항", 2)
restart_numbering(doc)
for text in [
    "UI/UX 동일성 요구를 프론트엔드 소스 코드 동결로 확대 해석한 근거가 종료보고서에 제시되어 있지 않다.",
    "프론트엔드 신규 개발 중단 및 기존 코드 보존이 범위 변경이라면, 변경요청·영향분석·승인·수정된 인수기준이 확인되어야 한다.",
    "종료보고서의 대표 성능 수치는 진단 생성, 목록 조회, 웹공시 응답 등 주로 백엔드·데이터 구조 개선 결과다. 프론트엔드 코드가 그대로라면 전체 사용자 체감 성능 개선을 입증하는 별도의 종단간(E2E) 지표가 필요하다.",
    "Django REST Framework(DRF)에서 Django Ninja로의 전환은 백엔드 프레임워크 교체이며, API 계약 보존의 필요성은 설명하지만 프론트엔드 수정 금지의 필연성을 만들지는 않는다.",
]:
    add_number(doc, text)

add_heading(doc, "1.3 요청 결론", 2)
add_text(
    doc,
    "종료보고서는 현 상태만으로 ‘프론트엔드 코드 보존이 UI/UX 동일성 요구의 당연한 귀결’이라는 인상을 줄 수 있으므로, 이를 확정 사실로 표현하기 전에 계약 기준선과 전략 변경 승인 자료를 확인하고, 백엔드 성능 개선과 전체 서비스 성능 개선을 구분하여 기술할 것을 요청한다."
)

add_page_break(doc)

# 2
add_heading(doc, "2. 종료보고서 기재 내용과 쟁점", 1)
add_heading(doc, "2.1 보고서가 스스로 밝힌 사실", 2)
add_table(
    doc,
    ["보고서 기재 사항", "확인되는 사실", "검토상 의미"],
    [
        ("2.2 범위 및 원칙", "프로젝트 범위를 백엔드 전면 재작성과 운영 데이터 이관으로 기재", "계약 원문에서 프론트엔드가 제외됐는지 확인 필요"),
        ("2.2 UI/UX 완전 보존", "사용자가 체감하는 화면과 동작을 동일하게 유지", "결과·행동의 동등성 기준이며 소스 동결 문구는 아님"),
        ("3.1 추진 전략", "초기에는 프론트·백엔드 모두 신규 기술로 재작성", "당초 수행계획에 프론트 리뉴얼이 포함됐음을 시사"),
        ("3.1 전략 전환", "2026-04-25 신규 프론트 개발 중단, 기존 화면 보존 결정", "중요 범위 변경이므로 승인 및 영향분석 확인 필요"),
        ("2.4 기술 스택", "V1: Django 4.2 + DRF 3.14 / V2: Django 6.0 + Django Ninja 1.5", "백엔드는 사실상 신규 구현·호환 어댑터 구축"),
        ("4.4 수행 규모", "중단된 신규 프론트 구현분 약 14.7만 라인 별도 표기", "상당한 수행분 폐기·전환에 대한 의사결정 근거 필요"),
    ],
    [4.0, 5.7, 6.3],
)

add_heading(doc, "2.2 동의하는 부분", 2)
for text in [
    "운영 중 서비스의 화면과 업무 흐름을 유지해 고객 재교육과 전환 위험을 줄이려는 목표는 합리적이다.",
    "신규 백엔드가 기존 API 경로·응답·인증 흐름을 호환하면 기존 프론트엔드를 연결해 전환 위험을 낮출 수 있다.",
    "데이터 구조, 쿼리, 캐시, 정적 스냅샷 등 백엔드 중심 개선만으로도 서버 처리량과 확장성을 크게 높일 수 있다.",
]:
    add_bullet(doc, text)

add_heading(doc, "2.3 동의하기 어려운 확대 해석", 2)
add_callout(
    doc,
    "구분이 필요한 두 문장",
    "① “사용자가 보는 UI/UX는 동일해야 한다.”  ② “기존 프론트엔드 소스와 빌드 산출물을 수정하지 않는다.” "
    "두 문장은 서로 다른 요구사항이다. ②가 요구되려면 코드 동결, 배포 산출물 유지, 변경 허용 범위가 계약·요구사항·변경승인 문서에 별도로 명시되어야 한다.",
    fill="FFF2CC",
    accent="BF9000",
)

add_page_break(doc)

# 3
add_heading(doc, "3. 기술적 검토", 1)
add_heading(doc, "3.1 UI/UX 동일성과 프론트엔드 코드 변경은 양립 가능", 2)
add_text(
    doc,
    "UI/UX는 외부에서 관찰되는 화면 구조, 시각적 표현, 입력·이동·오류 처리, 접근성 및 업무 흐름에 관한 결과 기준이다. 반면 소스 코드 동결은 내부 구현에 대한 제약이다. "
    "동일한 UI/UX를 유지하면서도 다음과 같은 프론트엔드 개선은 통상 가능하다."
)
for text in [
    "불필요한 재렌더링 제거, 메모이제이션 및 상태 구독 범위 축소",
    "코드 분할, 지연 로딩, 번들 트리 셰이킹 및 의존성 정리",
    "중복 API 호출 제거, 요청 병합·취소·캐시·프리패치 적용",
    "이미지·폰트·정적 자산 최적화와 CDN 캐시 정책 개선",
    "목록 가상화, 대용량 화면의 점진적 렌더링 및 메모리 누수 제거",
    "접근성·키보드 동작·오류 처리 유지 여부를 회귀 테스트로 보증",
]:
    add_bullet(doc, text)

add_heading(doc, "3.2 DRF → Django Ninja 전환의 의미", 2)
add_table(
    doc,
    ["구분", "V1", "V2", "프론트엔드에 대한 함의"],
    [
        ("백엔드 프레임워크", "Django REST Framework", "Django Ninja", "서버 내부 구현 교체"),
        ("스키마·검증", "Serializer 중심", "타입·스키마 중심", "응답 계약을 같게 만들 수 있으나 자동으로 같아지지는 않음"),
        ("라우팅·오류", "DRF 관례", "Ninja 관례", "호환 계층 및 계약 테스트 필요"),
        ("문서화", "OpenAPI 지원", "OpenAPI 자동화 강점", "프론트 타입 생성·계약 검증에 활용 가능"),
        ("성능", "구현·쿼리에 좌우", "낮은 오버헤드 기대 가능", "사용자 체감 성능은 네트워크·렌더링까지 함께 측정해야 함"),
    ],
    [3.1, 3.5, 3.5, 5.9],
)
add_text(
    doc,
    "즉, DRF에서 Django Ninja로의 변경은 백엔드 리뉴얼의 실체를 뒷받침한다. 그러나 기존 API 계약을 유지할 수 있다는 사실은 ‘기존 프론트도 연결 가능하다’는 선택지를 제공할 뿐, ‘프론트 코드는 수정해서는 안 된다’는 결론을 만들지 않는다."
)

add_heading(doc, "3.3 API 호환성과 프론트엔드 리뉴얼은 별개 축", 2)
add_text(
    doc,
    "API 호환성은 기존 클라이언트의 연속성을 위한 안전장치다. 프론트엔드 V2를 새로 만들거나 내부 구조를 개선하더라도 같은 API를 사용할 수 있다. 반대로 기존 프론트엔드를 유지하더라도 API 중복 호출, 렌더링 병목, 번들 비대화가 남아 있으면 사용자 체감 성능은 개선되지 않을 수 있다."
)

add_page_break(doc)

# 4
add_heading(doc, "4. 성능개선 주장에 대한 검증 범위", 1)
add_heading(doc, "4.1 종료보고서가 제시한 대표 지표", 2)
add_table(
    doc,
    ["지표", "보고서 수치", "주된 개선 계층", "판단 가능 범위"],
    [
        ("진단 생성", "366ms → 35ms, 10.5배", "DB 모델·쿼리·백엔드", "서버 처리 성능 개선"),
        ("목록 조회", "651ms → 55ms, 11.9배", "DB 읽기·스냅샷", "API 처리 성능 개선"),
        ("웹공시 첫 응답", "0.37초 → 0.04초, 약 10배", "정적 생성·CDN 캐시", "해당 응답 경로 개선"),
        ("전송 데이터", "159KB → 96KB, 약 40% 감소", "응답·정적 자산", "해당 시나리오 전송량 개선"),
    ],
    [3.6, 3.5, 4.2, 4.7],
)
add_text(
    doc,
    "위 수치는 백엔드 및 데이터 제공 계층의 개선을 보여주는 유의미한 근거다. 다만 전체 서비스의 프론트엔드 성능이 개선되었다는 결론을 내리려면 브라우저에서 사용자가 실제로 경험하는 종단간 지표가 추가되어야 한다."
)

add_heading(doc, "4.2 추가로 필요한 종단간 지표", 2)
add_table(
    doc,
    ["검증 영역", "권장 지표", "동일 조건"],
    [
        ("초기 로딩", "LCP, FCP, TTFB, 초기 JS/CSS 전송량", "동일 기기·망·캐시 상태·데이터"),
        ("상호작용", "INP, 입력 지연, 화면 전환 완료 시간", "동일 사용자 시나리오"),
        ("화면 안정성", "CLS, 레이아웃 이동 및 깜빡임", "동일 해상도·브라우저"),
        ("API 사용", "요청 수, 중복 요청, 오류율, 취소되지 않은 요청", "동일 업무 흐름"),
        ("자원 사용", "메모리, CPU, 장시간 사용 시 누수", "동일 세션 길이"),
        ("회귀", "스크린샷·DOM·접근성·업무 시나리오", "V1 기준선 대비"),
    ],
    [3.2, 7.0, 5.8],
)

add_heading(doc, "4.3 표현상 수정이 필요한 부분", 2)
add_text(
    doc,
    "종단간 증거가 없다면 보고서의 결론은 “백엔드 처리 구조 및 특정 제공 경로의 성능이 개선되었다”로 한정하는 것이 타당하다. "
    "“시스템 전체 성능” 또는 “사용자 체감 성능” 개선으로 표현하려면 위 브라우저 지표와 실제 업무 시나리오 결과가 함께 제시되어야 한다."
)

add_page_break(doc)

# 5
add_heading(doc, "5. 범위 변경 및 수용 기준 검토", 1)
add_heading(doc, "5.1 반드시 확인할 증빙", 2)
add_table(
    doc,
    ["구분", "확인 문서·증거", "확인할 핵심 질문"],
    [
        ("계약 기준선", "계약서, 제안서, 과업지시서, 요구사항 명세", "프론트·백엔드 V2 리뉴얼이 모두 범위였는가?"),
        ("전략 변경", "2026-04-25 결정 문서, 회의록, 변경요청서", "프론트 신규 개발 중단을 누가 어떤 조건으로 승인했는가?"),
        ("영향분석", "일정·비용·산출물·기술부채 분석", "14.7만 라인 중단분과 기존 코드 유지 리스크를 어떻게 처리했는가?"),
        ("수정 인수기준", "변경된 검수표와 승인 기록", "프론트 리뉴얼 대신 어떤 결과를 납품 완료로 인정했는가?"),
        ("성능 기준선", "측정 스크립트·원시 결과·환경 정보", "백엔드 지표와 종단간 지표가 분리되어 있는가?"),
        ("프론트 상태", "V1/V2 저장소 비교, 빌드·배포 이력", "‘그대로 보존’의 범위가 소스, 의존성, 설정, 배포 중 어디까지인가?"),
    ],
    [3.2, 6.0, 6.8],
)

add_heading(doc, "5.2 판단 기준", 2)
add_callout(
    doc,
    "판단 원칙",
    "명시적 승인과 수정된 인수기준이 확인되면 ‘기존 프론트 보존 + 백엔드 전면 재구축’은 합의된 범위 변경으로 평가할 수 있다. "
    "반대로 당초 범위가 프론트·백엔드 전체 리뉴얼인데 승인 자료가 없다면, UI/UX 동일성만을 이유로 프론트엔드 리뉴얼을 제외한 것은 과업 축소 또는 미이행 여부를 검토해야 한다.",
)

add_heading(doc, "5.3 권장 수용 기준", 2)
restart_numbering(doc)
for text in [
    "기능 동등성: 기존 업무 시나리오 전건 통과 및 회귀 이슈 종결",
    "시각 동등성: 주요 화면의 스크린샷 차이 기준과 허용 오차 합의",
    "API 계약: 경로·메서드·스키마·오류·인증에 대한 자동 계약 테스트",
    "백엔드 성능: 부하·처리량·지연·DB 사용량을 동일 환경에서 비교",
    "프론트 성능: 핵심 사용자 여정별 Core Web Vitals 및 자원 사용 비교",
    "운영성: 모니터링, 장애 대응, 롤백, 보안 및 유지보수 문서 인수",
]:
    add_number(doc, text)

add_page_break(doc)

# 6
add_heading(doc, "6. 공식 요청사항", 1)
add_text(doc, "본 이의제기에 따라 다음 조치를 요청한다.")
requests = [
    ("요구사항 기준선 제시", "프론트·백엔드 리뉴얼 및 성능개선 범위를 확인할 수 있는 계약·요구사항 문서를 제시한다."),
    ("전략 변경 승인 확인", "2026-04-25 프론트 신규 개발 중단의 요청자, 승인자, 승인 조건과 영향을 문서로 확인한다."),
    ("보고서 문구 정정", "UI/UX 동일성 요구와 프론트 코드 보존 결정을 인과적으로 동일시하지 않고, 별도 의사결정임을 명시한다."),
    ("성능 범위 분리", "백엔드·DB·CDN 개선과 브라우저 종단간 성능을 구분하여 성과를 기재한다."),
    ("프론트엔드 검증 보완", "프론트 코드를 유지했다면 기술부채·의존성·보안·번들·렌더링 성능을 진단하고 결과를 첨부한다."),
    ("인수 기준 재확인", "범위 변경이 승인된 경우 수정된 산출물과 검수 기준을 양측이 재확인한다."),
]
add_table(doc, ["요청", "세부 내용"], requests, [4.3, 11.7])

add_heading(doc, "6.1 요청하는 정정 문안", 2)
add_callout(
    doc,
    "권장 문안",
    "“UI/UX 동일성은 사용자에게 보이는 화면과 업무 흐름의 동등성을 의미한다. 프론트엔드 기존 코드 보존은 이 요구의 자동적 귀결이 아니라, 전환 위험·일정·비용·기술 검증 결과를 고려해 별도로 채택한 구현 전략이다. "
    "해당 전략의 적용 범위와 인수기준은 2026년 4월 25일 의사결정 자료 및 관련 승인 기록에 따른다.”",
    fill=PALE_BLUE,
)

add_heading(doc, "6.2 최종 의견", 2)
add_text(
    doc,
    "DRF 기반 V1 백엔드를 Django Ninja 기반 V2로 교체하고 기존 API 계약을 재현한 것은 분명한 백엔드 리뉴얼이다. 그러나 그것만으로 프론트엔드 V2 리뉴얼 요구가 충족되었다고 볼 수는 없다. "
    "UI/UX 동일성은 프론트 코드 개선과 양립 가능하며, 프론트 코드를 일절 수정하지 않는 조건은 별도의 명시적 요구 또는 승인된 범위 변경이 있어야 성립한다."
)
add_text(
    doc,
    "따라서 본 건의 적정성 판단은 ‘기술적으로 기존 프론트를 재사용할 수 있었는가’가 아니라, ‘당초 범위와 변경 절차에 따라 그 선택이 승인되었는가, 그리고 성능개선 성과가 합의된 검수 범위로 입증되었는가’를 기준으로 이루어져야 한다.",
    bold=True,
)

add_page_break(doc)

# Appendix
add_heading(doc, "부록 A. 참조 근거 매핑", 1)
add_table(
    doc,
    ["본 의견서 논점", "종료보고서 참조 위치", "참조 내용 요약"],
    [
        ("프로젝트 범위", "2.2 범위 및 원칙", "백엔드 전면 재작성과 데이터 이관"),
        ("UI/UX 기준", "2.2 범위 및 원칙", "사용자가 체감하는 화면·동작 동일"),
        ("초기 계획", "3.1 추진 전략", "프론트·백엔드 모두 신규 기술로 재작성"),
        ("전략 변경", "3.1 및 3.2", "2026-04-25 프론트 신규 개발 중단"),
        ("백엔드 기술 전환", "2.4 기술 스택 변화", "Django+DRF → Django+Django Ninja"),
        ("프론트 중단 규모", "3.1 및 4.4", "신규 프론트 약 14.7만 라인 구현 후 중단"),
        ("성능 수치", "5.1 성능 개선", "진단 생성·목록·웹공시 중심의 실측"),
        ("검증 결과", "6장 품질 및 검증", "API·화면·보안 시나리오 검증"),
    ],
    [4.0, 4.3, 7.7],
)

add_heading(doc, "부록 B. 유의사항", 1)
for text in [
    "본 의견서는 제공된 종료보고서 초안만을 근거로 하며, 계약 위반이나 법적 책임을 단정하지 않는다.",
    "최종 판단 전 계약서, 과업지시서, 제안서, 회의록, 변경승인서 및 검수기준을 함께 대조해야 한다.",
    "문서 내 미기입된 작성자·수신자·배포 구분은 제출 전에 확정해야 한다.",
]:
    add_bullet(doc, text)

doc.core_properties.title = "EDK 리뉴얼 프로젝트 이의제기 및 검토의견서"
doc.core_properties.subject = "UI/UX 동일성, 프론트엔드 코드 보존 및 성능개선 범위 검토"
doc.core_properties.author = ""
doc.core_properties.keywords = "EDK, 리뉴얼, 이의제기, DRF, Django Ninja, UI/UX, 성능개선"
settings = doc.settings.element
update_fields = settings.find(qn("w:updateFields"))
if update_fields is None:
    update_fields = OxmlElement("w:updateFields")
    settings.append(update_fields)
update_fields.set(qn("w:val"), "true")

doc.save(OUTPUT)
print(OUTPUT)
